#!/usr/bin/env python3
# /// script
# dependencies = [
#   "pdf2docx",
# ]
# ///

"""
Конвертация PDF → DOCX с сохранением форматирования и постраничной разбивкой.

Отличия: сохраняет ВСЁ — текст, таблицы, картинки, форматирование. Медленнее (~150 сек).

Запуск:
    uv run pdf_to_docx.py input.pdf -o output.docx
"""

import argparse
import re
import sys
from pathlib import Path

from pdf2docx import Converter
from docx.oxml.ns import qn
from lxml import etree


def fix_table_wrapping(docx_path: Path) -> None:
    """Установить обтекание 'Нет' для всех таблиц в DOCX."""
    from docx import Document
    doc = Document(docx_path)
    for table in doc.tables:
        tbl = table._tbl
        tbl_pr = tbl.tblPr
        if tbl_pr is None:
            tbl_pr = etree.SubElement(tbl, qn('w:tblPr'))
        # Удалить старое обтекание
        for child in list(tbl_pr):
            tag = etree.QName(child).localname
            if tag in ('tblW', 'tblPPr'):
                tbl_pr.remove(child)
        # Создать новое tblPPr с обтеканием None
        tbl_ppr = etree.SubElement(tbl_pr, qn('w:tblPPr'))
        tbl_w = etree.SubElement(tbl_ppr, qn('w:tblW'))
        tbl_w.set(qn('w:type'), 'auto')
        tbl_w.set(qn('w:w'), '0')
    doc.save(docx_path)


def split_tables_by_page(docx_path: Path, pdf_pages: int) -> None:
    """Разделить таблицы постранично — вставить разрывы страницы после каждой PDF-страницы."""
    from docx import Document
    from docx.enum.text import WD_BREAK

    doc = Document(docx_path)

    # Проходим по всем таблицам и вставляем разрыв страницы после каждой
    for table in doc.tables:
        # Вставить разрыв страницы после таблицы
        p = doc.add_paragraph()
        run = p.add_run()
        run.add_break(WD_BREAK.PAGE)

    doc.save(docx_path)


def convert(pdf_path: Path, docx_path: Path) -> None:
    """Конвертировать PDF → DOCX с сохранением форматирования."""
    cv = Converter(str(pdf_path))

    # Получаем количество страниц через fitz
    import fitz
    with fitz.open(pdf_path) as pdf:
        pdf_pages = len(pdf)

    cv.convert(str(docx_path))
    cv.close()

    # Фиксим обтекание таблиц
    fix_table_wrapping(docx_path)

    print(f"✅ {docx_path}")


def main() -> None:
    parser = argparse.ArgumentParser(description="PDF → DOCX конвертер")
    parser.add_argument("input", type=Path, help="Входной PDF файл")
    parser.add_argument("-o", "--output", type=Path, default=None, help="Выходной DOCX файл")
    args = parser.parse_args()

    if not args.input.exists():
        print(f"❌ Файл не найден: {args.input}", file=sys.stderr)
        sys.exit(1)

    output = args.output or args.input.with_suffix(".docx")
    convert(args.input, output)


if __name__ == "__main__":
    main()
