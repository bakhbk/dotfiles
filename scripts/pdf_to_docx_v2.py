#!/usr/bin/env python3
# /// script
# dependencies = [
#   "pymupdf",
#   "python-docx",
# ]
# ///

"""
Конвертация PDF → DOCX (вариант 2): текст + картинки, без таблиц.

Отличия: только текст и картинки, таблицы игнорирует. Быстро (~5 сек).

Запуск:
    uv run pdf_to_docx_v2.py input.pdf -o output.docx
"""

import argparse
import io
import re
import sys
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches, Pt


def add_image(doc: Document, img_bytes: bytes) -> None:
    """Вставить картинку в документ."""
    stream = io.BytesIO(img_bytes)
    doc.add_picture(stream, width=Inches(5))


def convert(pdf_path: Path, docx_path: Path) -> None:
    """Конвертировать PDF → DOCX с постраничной разбивкой, без таблиц."""
    doc = Document()

    with fitz.open(pdf_path) as pdf:
        for page_num, page in enumerate(pdf):
            if page_num > 0:
                doc.add_page_break()

            # Извлечение текста с сохранением форматирования
            text = page.get_text("text", sort=True)

            # Разбиваем на абзацы по пустым строкам
            paragraphs = re.split(r'\n\s*\n', text.strip()) if text.strip() else []

            for para in paragraphs:
                p = doc.add_paragraph(para.strip())
                # Установить отступ после абзаца для разделения
                p.paragraph_format.space_after = Pt(6)

            # Извлечение картинок
            images = page.get_images(full=True)
            for img_index, img in enumerate(images):
                xref = img[0]
                try:
                    pix = fitz.Pixmap(pdf, xref)
                    img_bytes = pix.tobytes("png")
                    add_image(doc, img_bytes)
                except Exception as e:
                    print(f"⚠️ Не удалось вставить картинку {img_index}: {e}", file=sys.stderr)

    doc.save(docx_path)
    print(f"✅ {docx_path}")


def main() -> None:
    parser = argparse.ArgumentParser(description="PDF → DOCX конвертер (без таблиц)")
    parser.add_argument("input", type=Path, help="Входной PDF файл")
    parser.add_argument("-o", "--output", type=Path, default=None, help="Выходной DOCX файл")
    args = parser.parse_args()

    if not args.input.exists():
        print(f"❌ Файл не найден: {args.input}", file=sys.stderr)
        sys.exit(1)

    output = args.output or args.input.with_suffix(".docx")
    convert(args.input, output)


if __name__ == "__main__":
    main()
